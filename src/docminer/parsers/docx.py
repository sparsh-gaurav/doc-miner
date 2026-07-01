"""Text extraction for Word (.docx) documents."""

from docx import Document


def extract_text(path: str) -> str:
    """Extract paragraph and table text from a .docx file."""
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]

    # also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    text = "\n".join(lines)
    if not text.strip():
        raise ValueError("No text extracted from Word document.")
    return text
